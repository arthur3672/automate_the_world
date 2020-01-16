#!/usr/bin/python3

import xml.etree.ElementTree as ET
import string
import os
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
import argparse

def get_ready(output_file):
	document = Document()
	document.add_heading('Service Detection')
	document.save(output_file)

def main(nmap_file, output_file):
	tree = ET.parse(nmap_file)
	root = tree.getroot()

	document = Document(output_file)

	total_host = len(root)
	for hosts in range(0, total_host):
		if root[hosts].tag == 'host' and root[hosts][0].attrib['state'] == 'up':
			ip_addr = root[hosts][1].attrib['addr']
			find_port_element = int(len(root[hosts]))
			for ports in range (2, find_port_element):
				if root[hosts][ports].tag == 'ports':
					num_port = int(len(root[hosts][ports]))
					count = 1
					document.add_heading(ip_addr, level=2)
					table = document.add_table(rows=1, cols=3)
					table.style = 'Table Grid'
					table.cell(0,0).text = 'Port'
					table.cell(0,1).text = 'Service Name'
					table.cell(0,2).text = 'Version'
					for i in range(0, num_port):
						if str(root[hosts][ports][i].tag) == 'port':
							if str(root[hosts][ports][i][0].attrib['state']) == 'open':
								port = str(root[hosts][ports][i].attrib['portid'])+'/'+str(root[hosts][ports][i].attrib['protocol'])
								if len(root[hosts][ports][i]) > 1:
									if 'name' in str(root[hosts][ports][i][1].attrib.keys()):
										service_name = str(root[hosts][ports][i][1].attrib['name'])
									else:
										service_name = 'unknown'
									if 'product' in str(root[hosts][ports][i][1].attrib.keys()):
										product = str(root[hosts][ports][i][1].attrib['product'])
									else:
										product = ''
									if 'version' in str(root[hosts][ports][i][1].attrib.keys()):
										version = str(root[hosts][ports][i][1].attrib['version'])
									else:
										version = ''
									if 'extrainfo' in str(root[hosts][ports][i][1].attrib.keys()):
										extrainfo = str(root[hosts][ports][i][1].attrib['extrainfo'])
									else:
										extrainfo = ''
									versions = product+' '+version+' '+extrainfo
									versions = versions.strip()
									if versions == '':
										versions = '-'
								else:
									service_name = 'unknown'
									versions = '-'

								table.add_row()
								table.cell(count,0).text = port
								table.cell(count,1).text = service_name
								table.cell(count,2).text = versions
								count += 1
					
					widths = (Inches(1), Inches(2), Inches(3))
					for row in table.rows:
						for idx, width in enumerate(widths):
							row.cells[idx].width = width
					document.add_paragraph('')
					document.save(output_file)
					break
	document.add_page_break()
	document.save(output_file)

if __name__ == '__main__':
	parser = argparse.ArgumentParser()
	parser.add_argument('-f', metavar='Nmap XML File / Directory', type=str, help='Nmap XML file / directory', required=True)
	parser.add_argument('-o', metavar='Output file', type=str, help='Output DOCX file', required=True)
	args = parser.parse_args()
	
	if not args.o.endswith('.docx'):
		output_file = args.o + '.docx'
	else:
		output_file = args.o
	get_ready(output_file)

	if os.path.isdir(args.f):
		for filename in os.listdir(args.f):
			if filename.endswith('.xml'):
				fullpath_file = args.f + filename
				main(fullpath_file, output_file)
	else:
		if args.f.endswith('.xml'):
			main(args.f, output_file)
		else:
			parser.print_usage()
	exit()