#!/usr/bin/python3

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl import load_workbook
import cis_raw_to_xlsx_os
import string
import argparse
import os

def main(input_file, output_file):
	input_xlsx_file = load_workbook(filename=input_file, read_only=True)
	document = Document()
	document.add_heading('Findings Summary')
	document.save(output_file)
	for x in input_xlsx_file.sheetnames:
		ws = input_xlsx_file[x]
		row_count = ws.max_row + 1

		document = Document(output_file)
		document.add_heading(x, level=2)
		red_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
		red_2 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))

		current_row = 2
		table = document.add_table(rows=1, cols=2)
		table.style = 'Table Grid'
		table.rows[0].cells[0]._tc.get_or_add_tcPr().append(red_1)
		table.rows[0].cells[1]._tc.get_or_add_tcPr().append(red_2)
		table.cell(0, 0).text = 'No'
		table.cell(0, 1).text = 'Findings'
		count = 1
		for i in range(current_row, row_count):
			table.add_row()
			table.cell(count, 0).text = str(count)
			table.cell(count, 1).text = ws.cell(row=i, column=2).value
			count += 1
		widths = (Inches(0.4), Inches(5.5))
		for row in table.rows:
			for idx, width in enumerate(widths):
				row.cells[idx].width = width
		document.add_page_break()
		document.save(output_file)

	document = Document(output_file)
	for x in input_xlsx_file.sheetnames:
		ws = input_xlsx_file[x]
		row_count = ws.max_row + 1
		document.add_heading('Finding Details' +' ('+x+')', level=1)
		document.save(output_file)
		for i in range(2, row_count):
			name = ws.cell(row=i, column=2).value
			description = ws.cell(row=i, column=3).value
			severity = ws.cell(row=i, column=4).value
			current_host_value = ws.cell(row=i, column=5).value
			affected_host = ws.cell(row=i, column=6).value
			remediation = ws.cell(row=i, column=7).value
			status = ws.cell(row=i, column=8).value
			document.add_heading(name, level=2)
			document.add_heading('Description', level=3)
			for desc in description.split('\n\n'):
				document.add_paragraph(desc)
			document.add_paragraph()
			document.add_heading('Risk Level', level=3)
			document.add_paragraph(severity)
			document.add_paragraph()
			document.add_heading('Current Host Value', level=3)
			for value in current_host_value.split('\r\n'):
				document.add_paragraph(value)
			document.add_paragraph()
			document.add_heading('Affected Host', level=3)
			for hosts in affected_host.split('\r\n'):
				document.add_paragraph(hosts)
			document.add_paragraph()
			document.add_heading('Remediation', level=3)
			for solution in remediation.split('\n\n'):
				document.add_paragraph(solution)
			document.add_paragraph()
			document.add_heading('Status', level=3)
			document.add_paragraph(status)
			document.add_page_break()
			document.save(output_file)
	print('Convert done! Documentation saved to '+output_file)

if __name__ == '__main__':
	parser = argparse.ArgumentParser()
	parser.add_argument('-f', metavar='Raw Nessus File / Directory', type=str, help='Raw Nessus File / Directory', required=True)
	parser.add_argument('-o', metavar='Output file', type=str, help='Output DOCX file', required=True)
	args = parser.parse_args()

	if args.o.endswith('.docx'):
		output_xlsx_file = args.o[:-5] + '.xlsx'
		workbook = Workbook()
		workbook.save(filename=output_xlsx_file)
		if os.path.isdir(args.f):
			for filename in os.listdir(args.f):
				if filename.endswith('.nessus'):
					fullpath_file = args.f + filename
					cis_raw_to_xlsx_os.main(fullpath_file, output_xlsx_file)
		else:
			cis_raw_to_xlsx_os.main(args.f, output_xlsx_file)
		
		workbook = load_workbook(output_xlsx_file)
		workbook.remove(workbook['Sheet'])
		workbook.save(filename=output_xlsx_file)

		main(output_xlsx_file, args.o)
	else:
		parser.print_usage()

	exit()