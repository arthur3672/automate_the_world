#!/usr/bin/python3

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches
import vapt_raw_to_xlsx
import vapt_nmap
import string
import argparse
import os
from openpyxl import load_workbook
from datetime import datetime

def main(input_file, output_file):
	input_xlsx_file = load_workbook(filename=input_file, read_only=True)
	ws = input_xlsx_file['Items']
	row_count = ws.max_row + 1

	document = Document(output_file)

	red_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
	red_2 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
	red_3 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
	orange_1 = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
	orange_2 = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
	orange_3 = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
	yellow_1 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
	yellow_2 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
	yellow_3 = parse_xml(r'<w:shd {} w:fill="FFFF00"/>'.format(nsdecls('w')))
	green_1 = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
	green_2 = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
	green_3 = parse_xml(r'<w:shd {} w:fill="00FF00"/>'.format(nsdecls('w')))
	blue_1 = parse_xml(r'<w:shd {} w:fill="00BFFF"/>'.format(nsdecls('w')))
	blue_2 = parse_xml(r'<w:shd {} w:fill="00BFFF"/>'.format(nsdecls('w')))
	blue_3 = parse_xml(r'<w:shd {} w:fill="00BFFF"/>'.format(nsdecls('w')))

	current_row = 2

	document.add_heading('Findings Summary')

	if float(ws.cell(row=current_row, column=6).value) >= 9:
		document.add_heading('Critical Findings', level=2)
		table = document.add_table(rows=1, cols=3)
		table.style = 'Table Grid'
		table.rows[0].cells[0]._tc.get_or_add_tcPr().append(red_1)
		table.rows[0].cells[1]._tc.get_or_add_tcPr().append(red_2)
		table.rows[0].cells[2]._tc.get_or_add_tcPr().append(red_3)
		table.cell(0, 0).text = 'No'
		table.cell(0, 1).text = 'Findings'
		table.cell(0, 2).text = 'Affected IP / URL / Host'
		count = 1
		for i in range(current_row, row_count):
			if float(ws.cell(row=i, column=6).value) < 9:
				break
			else:
				table.add_row()
				table.cell(count, 0).text = str(count)
				table.cell(count, 1).text = ws.cell(row=i, column=2).value
				if len(ws.cell(row=i, column=4).value.split('\r\n')) == 1:
					table.cell(count, 2).text = ws.cell(row=i, column=4).value
				else:
					for ip in ws.cell(row=i, column=4).value.split('\r\n'):
						table.cell(count, 2).text += ip+'\n'
					table.cell(count, 2).text = table.cell(count, 2).text[:-1]
				count += 1
		current_row = i
		widths = (Inches(0.4), Inches(3.7), Inches(2))
		for row in table.rows:
			for idx, width in enumerate(widths):
				row.cells[idx].width = width
		document.add_page_break()

	if float(ws.cell(row=current_row, column=6).value) >= 7:
		document.add_heading('High Findings', level=2)
		table = document.add_table(rows=1, cols=3)
		table.style = 'Table Grid'
		table.rows[0].cells[0]._tc.get_or_add_tcPr().append(orange_1)
		table.rows[0].cells[1]._tc.get_or_add_tcPr().append(orange_2)
		table.rows[0].cells[2]._tc.get_or_add_tcPr().append(orange_3)
		table.cell(0, 0).text = 'No'
		table.cell(0, 1).text = 'Findings'
		table.cell(0, 2).text = 'Affected IP / URL / Host'
		count = 1
		for i in range(current_row, row_count):
			if float(ws.cell(row=i, column=6).value) < 7:
				break
			else:
				table.add_row()
				table.cell(count, 0).text = str(count)
				table.cell(count, 1).text = ws.cell(row=i, column=2).value
				if len(ws.cell(row=i, column=4).value.split('\r\n')) == 1:
					table.cell(count, 2).text = ws.cell(row=i, column=4).value
				else:
					for ip in ws.cell(row=i, column=4).value.split('\r\n'):
						table.cell(count, 2).text += ip+'\n'
					table.cell(count, 2).text = table.cell(count, 2).text[:-1]
				count += 1
		current_row = i
		widths = (Inches(0.4), Inches(3.7), Inches(2))
		for row in table.rows:
			for idx, width in enumerate(widths):
				row.cells[idx].width = width
		document.add_page_break()

	if float(ws.cell(row=current_row, column=6).value) >= 4:
		document.add_heading('Medium Findings', level=2)
		table = document.add_table(rows=1, cols=3)
		table.style = 'Table Grid'
		table.rows[0].cells[0]._tc.get_or_add_tcPr().append(yellow_1)
		table.rows[0].cells[1]._tc.get_or_add_tcPr().append(yellow_2)
		table.rows[0].cells[2]._tc.get_or_add_tcPr().append(yellow_3)
		table.cell(0, 0).text = 'No'
		table.cell(0, 1).text = 'Findings'
		table.cell(0, 2).text = 'Affected IP / URL / Host'
		count = 1
		for i in range(current_row, row_count):
			if float(ws.cell(row=i, column=6).value) < 4:
				break
			else:
				table.add_row()
				table.cell(count, 0).text = str(count)
				table.cell(count, 1).text = ws.cell(row=i, column=2).value
				if len(ws.cell(row=i, column=4).value.split('\r\n')) == 1:
					table.cell(count, 2).text = ws.cell(row=i, column=4).value
				else:
					for ip in ws.cell(row=i, column=4).value.split('\r\n'):
						table.cell(count, 2).text += ip+'\n'
					table.cell(count, 2).text = table.cell(count, 2).text[:-1]
				count += 1
		current_row = i
		widths = (Inches(0.4), Inches(3.7), Inches(2))
		for row in table.rows:
			for idx, width in enumerate(widths):
				row.cells[idx].width = width
		document.add_page_break()

	if float(ws.cell(row=current_row, column=6).value) > 0:
		document.add_heading('Low Findings', level=2)
		table = document.add_table(rows=1, cols=3)
		table.style = 'Table Grid'
		table.rows[0].cells[0]._tc.get_or_add_tcPr().append(green_1)
		table.rows[0].cells[1]._tc.get_or_add_tcPr().append(green_2)
		table.rows[0].cells[2]._tc.get_or_add_tcPr().append(green_3)
		table.cell(0, 0).text = 'No'
		table.cell(0, 1).text = 'Findings'
		table.cell(0, 2).text = 'Affected IP / URL / Host'
		count = 1
		for i in range(current_row, row_count):
			if float(ws.cell(row=i, column=6).value) == 0:
				break
			else:
				table.add_row()
				table.cell(count, 0).text = str(count)
				table.cell(count, 1).text = ws.cell(row=i, column=2).value
				if len(ws.cell(row=i, column=4).value.split('\r\n')) == 1:
					table.cell(count, 2).text = ws.cell(row=i, column=4).value
				else:
					for ip in ws.cell(row=i, column=4).value.split('\r\n'):
						table.cell(count, 2).text += ip+'\n'
					table.cell(count, 2).text = table.cell(count, 2).text[:-1]
				count += 1
		current_row = i
		widths = (Inches(0.4), Inches(3.7), Inches(2))
		for row in table.rows:
			for idx, width in enumerate(widths):
				row.cells[idx].width = width
		document.add_page_break()

	if row_count - current_row > 1:
		document.add_heading('Info Findings', level=2)
		table = document.add_table(rows=1, cols=3)
		table.style = 'Table Grid'
		table.rows[0].cells[0]._tc.get_or_add_tcPr().append(blue_1)
		table.rows[0].cells[1]._tc.get_or_add_tcPr().append(blue_2)
		table.rows[0].cells[2]._tc.get_or_add_tcPr().append(blue_3)
		table.cell(0, 0).text = 'No'
		table.cell(0, 1).text = 'Findings'
		table.cell(0, 2).text = 'Affected IP / URL / Host'
		count = 1
		for i in range(current_row, row_count):
			table.add_row()
			table.cell(count, 0).text = str(count)
			table.cell(count, 1).text = ws.cell(row=i, column=2).value
			if len(ws.cell(row=i, column=4).value.split('\r\n')) == 1:
				table.cell(count, 2).text = ws.cell(row=i, column=4).value
			else:
				for ip in ws.cell(row=i, column=4).value.split('\r\n'):
					table.cell(count, 2).text += ip+'\n'
				table.cell(count, 2).text = table.cell(count, 2).text[:-1]
			count += 1
		widths = (Inches(0.4), Inches(3.7), Inches(2))
		for row in table.rows:
			for idx, width in enumerate(widths):
				row.cells[idx].width = width
		document.add_page_break()
	document.save(output_file)

	document.add_heading('Finding Details', level=1)
	for i in range(2, row_count):
		name = ws.cell(row=i, column=2).value
		description = ws.cell(row=i, column=3).value
		affected_host = ws.cell(row=i, column=4).value
		remediation = ws.cell(row=i, column=5).value
		cvss_score = 'CVSS Score: '+str(ws.cell(row=i, column=6).value)
		severity = ws.cell(row=i, column=7).value
		status = ws.cell(row=i, column=8).value

		document.add_heading(name, level=2)
		document.add_heading('Vulnerability Description', level=3)
		for desc in description.split('\n\n'):
			document.add_paragraph(desc)
		document.add_paragraph()
		document.add_heading('Risk Level', level=3)
		document.add_paragraph(severity)
		document.add_paragraph(cvss_score)
		document.add_paragraph()
		document.add_heading('Affected IP / URL / Host', level=3)
		for hosts in affected_host.split('\r\n'):
			document.add_paragraph(hosts)
		document.add_paragraph()
		document.add_heading('Remediation', level=3)
		for solution in remediation.split('\n\n'):
			document.add_paragraph(solution)
		document.add_paragraph()
		document.add_heading('Status', level=3)
		document.add_paragraph(status)
		document.add_paragraph()
		document.add_heading('Proof of Concept', level=3)
		num = 9
		while isinstance(ws.cell(row=i, column=num).value, str):
			plugin_output = ws.cell(row=i, column=num).value
			for output in plugin_output.split('\r\n'):
				if '\n ' in output:
					output = output.replace('\n ', '', 1)
				if '\n ' in output:
					output = output.replace('\n ', '\n')
				document.add_paragraph(output)
			num += 1
		document.add_page_break()
		document.save(output_file)

	print('Convert done! Documentation saved to '+output_file)

if __name__ == '__main__':
	start_time = datetime.now()
	parser = argparse.ArgumentParser()
	parser.add_argument('-n', metavar='Raw Nessus File / Directory', type=str, help='Raw Nessus File / Directory', required=True)
	parser.add_argument('-x', metavar='Raw Nmap XML File / Directory', type=str, help='Raw Nmap XML File / Directory', required=True)
	parser.add_argument('-o', metavar='Output file', type=str, help='Output DOCX file', required=True)
	args = parser.parse_args()

	if args.o.endswith('.docx'):
		output_file = args.o
	else:
		parser.print_usage()
		exit()

	mid_xlsx_file = ''
	output_xlsx_file = output_file[:-5] + '.xlsx'
	if '/' in output_xlsx_file:
		mid_xlsx_file = output_xlsx_file.split('/')
		mid_xlsx_file[-1] = 'mid_'+mid_xlsx_file[-1]
		mid_xlsx_file = '/'.join(mid_xlsx_file)
	else:
		mid_xlsx_file = 'mid_'+output_xlsx_file
	vapt_raw_to_xlsx.preparation(mid_xlsx_file)

	if os.path.isdir(args.n):
		if args.n.endswith('/') == False:
			args.n += '/'
		for filename in os.listdir(args.n):
			if filename.endswith('.nessus'):
				fullpath_file = args.n + filename
				vapt_raw_to_xlsx.main(fullpath_file, mid_xlsx_file)
	else:
		if args.n.endswith('.nessus'):
			vapt_raw_to_xlsx.main(args.n, mid_xlsx_file)
		else:
			parser.print_usage()
			exit()

	os.remove(mid_xlsx_file)
	current_time = datetime.now()
	used_time = current_time - start_time
	print('Excel Completed! Time used: '+str(used_time))
	vapt_nmap.get_ready(output_file)

	if os.path.isdir(args.x):
		if args.x.endswith('/') == False:
			args.x += '/'
		for filename in os.listdir(args.x):
			if filename.endswith('.xml'):
				fullpath_file = args.x + filename
				vapt_nmap.main(fullpath_file, output_file)
	else:
		if args.x.endswith('.xml'):
			vapt_nmap.main(args.x, output_file)
		else:
			parser.print_usage()
			exit()

	main(output_xlsx_file, output_file)
	current_time = datetime.now()
	used_time = current_time - start_time
	print('Total Time used: '+str(used_time))
	exit()